from src.loaders.object import Object
from src.loaders.file import File
from src.loaders.tags import load_tags, Tags

import os
import re
import pandas as pd
import docx.table  # docx is in python-docx
from win32com import client
import dateparser
from dateparser.search import search_dates

from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph

above_filters = ['as above', 'above table']
below_filters = ['as below', 'below table']

def doc2docx(file_path: str) -> str:  # TODO: add try-except
    """
    convert doc to docx
    Args:
        file_path: file (path) to convert，due to pywin32，use absolute path

    Returns:
        converted docx file path
    """
    a = os.path.split(file_path)
    b = os.path.splitext(a[-1])[0]
    new_file_path = "{}\\{}.docx".format(a[0], b)
    print(f"doc file {b} detected，converting to docx")

    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(file_path)
    doc.SaveAs(new_file_path, 12)
    doc.Close()
    word.Quit()

    return new_file_path


def is_pure_number(s: str) -> bool:
    """
    Check if string is made of pure numbers
    Args:
        s: String to check

    Returns:
        True if pure number or space，False otherwise
    """
    if s == '':
        return True
    try:
        float(s)
        return True
    except ValueError:
        return False


def split_tables_by_numbers(table: list[list]) -> list[list[list]]:
    """
    Split multiple tables that are in one word table
    Args:
        table: list[list] table，by table_reader

    Returns:
        list[list[list]]，each detected header for one list
    """
    split_tables = []
    current_table = []

    for row in table:
        contains_pure_number = any(is_pure_number(cell) for cell in row)
        if not contains_pure_number:
            # No pure number, a new table, split
            if current_table:
                split_tables.append(current_table)
                current_table = []
        current_table.append(row)

    # last table
    if current_table:
        split_tables.append(current_table)

    return split_tables


# def find_context_tags(texts: list, tags: set) -> list:
#     tags_in_context = set()
#     for text in texts:
#         keyword_tags = set([tag for tag in tags if tag in text])
#         tags_in_context = tags_in_context | keyword_tags
#         time_tags = find_time_tag(text)
#         if time_tags != "":
#             tags_in_context.add(time_tags)
#
#     return list(tags_in_context)


def add_table_content_tags(table: Object, tags: Tags):
    """
    Add tags to table
    Args:
        table: table Object needed for tagging
        tags: Tags
    """
    tags_in_table = set()  # No date
    for column in table.content.values():
        col = list(column.values())
        for cell in col:
            if not is_pure_number(cell):
                tags.add_tag(cell, table.object_id)
                tags_in_table.add(cell)
            # for tag in tags.tags_dict:
            #     if tag in cell:
            #         tags.add_tag(tag, table.object_id)
            #         tags_in_table.add(tag)
    table.tags = tags_in_table


def add_text_content_tags(text: Object, tags: Tags):
    """
    Add tags to texts
    Args:
        text: text Object needed for tagging
        tags: Tags
    """
    tags_in_text = set()
    for tag in tags.tags_dict:
        if tag in text.content:
            tags.add_tag(tag, text.object_id)
            tags_in_text.add(tag)
    text.tags = tags_in_text

def add_date_tags(object: Object, tags: Tags, content: str):
    """
    Add date tags to Object
    Args:
        object: Object
        tags: Tags
        content: content that needs tagging
    """
    file_path = os.path.split(object.file_name)
    title = os.path.splitext(file_path[-1])[0].replace("_", " ")

    time_tag_title = None
    time_tag_content = None

    result_title = search_dates(title, settings={"PREFER_DAY_OF_MONTH": "first"})
    if result_title:
        for _, dt in result_title:
            time_tag_title = dt.strftime("%B %Y")

    result_content = search_dates(content, settings={"PREFER_DAY_OF_MONTH": "first"})
    if  result_content:
        for _, dt in result_content:
            time_tag_content = dt.strftime("%B %Y")

    if result_content:
        tags.add_tag(time_tag_content, object.object_id)
        object.date.add(time_tag_content)
    elif result_title:
        tags.add_tag(time_tag_title, object.object_id)
        object.date.add(time_tag_title)


def find_table_contexts(texts, table_num, above_range=1, below_range=1):
    """
    Find table contexts，unit is paragraph
    Args:
        texts: all paragraphs
        table_num: table id
        above_range: search range above
        below_range: search range below

    Returns:
        Context
    """
    above_range = max(table_num - above_range, 0)
    below_range = min(table_num + below_range, len(texts) - 1)

    above_contexts = texts[above_range: table_num]
    below_contexts = texts[table_num + 1: below_range + 1]

    # Remove unrelated contexts(like 'above table' in context below)
    filtered_above_contexts = []
    for context in above_contexts:
        if not [tag for tag in above_filters if tag in context]:
            filtered_above_contexts.append(context)
    # print("Above：", filtered_above_contexts)

    filtered_below_contexts = []
    for context in below_contexts:
        if not [tag for tag in below_filters if tag in context]:
            filtered_below_contexts.append(context)
    # print("Below：", filtered_below_contexts)
    return filtered_above_contexts, filtered_below_contexts


def find_text_contexts(texts, text_num, above_range=1, below_range=1):
    """
    Find text contexts，unit is paragraph
    Args:
        texts: all paragraphs
        text_num: text id
        above_range: search range above
        below_range: search range below

    Returns:
        Context
    """
    above_range = max(text_num - above_range, 0)
    below_range = min(text_num + below_range, len(texts) - 1)

    above_contexts = texts[above_range: text_num]
    below_contexts = texts[text_num + 1: below_range + 1]

    return above_contexts, below_contexts


def table_reader(table: docx.table.Table) -> list[list]:
    """
    Read table blocks, remove bad characters and convert to lowercase
    Args:
        table: docx.table.Table

    Returns:
        list[list] table
    """
    table_list = []
    for i, row in enumerate(table.rows):
        row_content = []
        for cell in row.cells:  # all cells in a row
            c = cell.text
            c = c.replace('\r', '').replace('\x07', '').replace('\n', '').strip()
            c = c.lower()
            row_content.append(c)
        table_list.append(row_content)
    table_list = split_tables_by_numbers(table_list)
    return table_list


def iter_block_items(parent):
    """
    Read 'blocks' in docx, from https://blog.csdn.net/m0_51944012/article/details/141025820
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("Error Processing Word")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def find_title(texts: list) -> str:
    """
    Find table name in above context
    """
    if len(texts) >= 1:
        return texts[-1].content.split('，')[-1].replace('：', '')
    else:
        return ""


def add_tags_and_contexts(tables, texts, elements, tags: Tags):
    """
    Add tags and context for tables and texts
    Args:
        tables: list of tables
        texts: list of texts
        elements: all elements
        tags: Tags
    """
    for table in tables:
        tb_num = table.position
        above_context, below_context = find_table_contexts(elements, tb_num, 1, 1)
        table.above = above_context[0] if above_context else ""  # Assume range 1
        table.below = below_context[0] if below_context else ""
        add_table_content_tags(table, tags)
        add_date_tags(table, tags, table.title)

    for text in texts:
        tx_num = text.position
        above_context, below_context = find_text_contexts(elements, tx_num, 1, 1)
        text.above = above_context[0] if above_context else ""
        text.below = below_context[0] if below_context else ""
        add_text_content_tags(text, tags)
        add_date_tags(text, tags, text.content)


def doc2table(file: File, tags: Tags) -> tuple[list[Object], list]:
    """
    Scan all tables and texts in docx
    Args:
        file: docx File that needs processing
        tags: Tags

    Returns:
        (tables, texts)
    """
    file_path = file.file_path
    file_name = os.path.splitext(os.path.split(file_path)[-1])[0]
    if not os.path.exists(file_path):
        print(f"Bad file, skipping {file_path}")
        return [], []

    recycle_flag = False
    recycle_name = None
    if os.path.splitext(file_path)[1].lower() == '.doc':
        recycle_name = doc2docx(file_path)
        docx_file = docx.Document(recycle_name)
        recycle_flag = True
    else:
        docx_file = docx.Document(file_path)

    docx_all_texts = []
    docx_all_tables = []
    docx_all_elements = []

    block_idx = 0
    for block in iter_block_items(docx_file):
        if isinstance(block, Paragraph):  # Text
            if block.text == '':  # Skip empty
                continue
            text = block.text.replace('\t', '').strip()
            text_object = Object(file_id=file.file_id, file_name=file_name,
                                position=block_idx, title="", content=text)
            docx_all_texts.append(text_object)
            docx_all_elements.append(text)
            block_idx += 1

        elif isinstance(block, Table):  # Table
            title = find_title(docx_all_texts)

            tables = table_reader(block)
            for table in tables:
                if len(table) <= 1:
                    continue
                dataframe = pd.DataFrame(table[1:], columns=table[0])
                dict_table = dataframe.to_dict("dict")

                table = Object(file_id=file.file_id, file_name=file_name, position=block_idx,
                                     title=title, content=dict_table)
                docx_all_tables.append(table)
                docx_all_elements.append(f"<Table:{table.object_id}>")
                block_idx += 1

    add_tags_and_contexts(docx_all_tables, docx_all_texts, docx_all_elements, tags)

    if recycle_flag:
        os.remove(recycle_name)

    return docx_all_tables, docx_all_texts


if __name__ == "__main__":
    tags = load_tags()
    file = File(file_path=
                     "D:\\CS\\CS510\\final-project\\data\\inputdata\\Store\\January_2025_Sales_Report.docx")

    tbs, txs = doc2table(file, tags)
    for i in tbs:
        i.display()
        print("---------------------")

    # for j in tags:
    #     if j.related_object_ids:
    #         print(j.tag_name, j.related_object_ids)

    # for j in txs:
    #     j.display()
    #     print("---------------------")
